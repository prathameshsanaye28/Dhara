# # # from flask import Flask
# # # from api.smp_analyzer import smp_analyzer_blueprint
# # # from api.summarizer import summarizer_blueprint
# # # from api.weather import weather_blueprint
# # # from api.maintain import maintain_blueprint
# # # from api.segmentation import segment_blueprint

# # # app = Flask(__name__)

# # # # Register Blueprints
# # # app.register_blueprint(smp_analyzer_blueprint, url_prefix='/smp-analyzer')
# # # app.register_blueprint(summarizer_blueprint, url_prefix='/summarizer')
# # # app.register_blueprint(weather_blueprint, url_prefix='/weather')
# # # app.register_blueprint(maintain_blueprint, url_prefix='/maintain')
# # # app.register_blueprint(segment_blueprint, url_prefix='/segment')

# # # @app.route('/')
# # # def index():
# # #     return "Welcome to the Unified Flask API! Use /smp-analyzer, /summarizer, or /weather."

# # # if __name__ == '__main__':
# # #     app.run(debug=True)


# # import os
# # os.environ['TF_ENABLE_ONEDNN_OPTS'] = '0'  # Disable oneDNN custom operations

# # from flask import Flask, request, jsonify
# # import io
# # import PyPDF2
# # import docx
# # from transformers import pipeline

# # app = Flask(__name__)

# # def extract_text_from_pdf(file_object):
# #     """
# #     Extract text from a PDF file object.
# #     """
# #     try:
# #         pdf_reader = PyPDF2.PdfReader(file_object)
# #         text = ""
# #         for page in pdf_reader.pages:
# #             text += page.extract_text() + "\n"
# #         return text
# #     except Exception as e:
# #         print(f"Error extracting text from PDF: {e}")
# #         return ""

# # def extract_text_from_docx(file_object):
# #     """
# #     Extract text from a DOCX file object.
# #     """
# #     try:
# #         doc = docx.Document(file_object)
# #         text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
# #         return text
# #     except Exception as e:
# #         print(f"Error extracting text from DOCX: {e}")
# #         return ""

# # def summarize_text(full_text, max_length=500, min_length=200):
# #     """
# #     Summarize given text.
# #     """
# #     summarizer = pipeline("summarization", model="facebook/bart-large-cnn")
# #     summary = summarizer(full_text, max_length=max_length, min_length=min_length, do_sample=False)
# #     return summary[0]['summary_text']

# # @app.route('/summarize_documents', methods=['POST'])
# # def summarize_documents():
# #     """
# #     Handle document summarization.
# #     """
# #     if 'files' not in request.files:
# #         return jsonify({'error': 'No files part in the request'}), 400
    
# #     files = request.files.getlist('files')
    
# #     if not files or files[0].filename == '':
# #         return jsonify({'error': 'No files selected for uploading'}), 400
    
# #     try:
# #         full_text = ""
# #         processed_files = []
        
# #         for file in files:
# #             file_object = io.BytesIO(file.read())
            
# #             if file.filename.lower().endswith('.pdf'):
# #                 text = extract_text_from_pdf(file_object)
# #             elif file.filename.lower().endswith(('.docx', '.doc')):
# #                 text = extract_text_from_docx(file_object)
# #             else:
# #                 continue
            
# #             if text.strip():
# #                 full_text += text + "\n\n---Document Separator---\n\n"
# #                 processed_files.append(file.filename)
        
# #         if not full_text.strip():
# #             return jsonify({'error': 'No text could be extracted from the uploaded files'}), 400
        
# #         summary = summarize_text(full_text)
        
# #         return jsonify({
# #             'summary': summary
# #         })
    
# #     except Exception as e:
# #         return jsonify({'error': str(e)}), 500

# # if __name__ == '__main__':
# #     app.run(debug=False, port=3000, host='0.0.0.0')


# import os
# os.environ['TF_ENABLE_ONEDNN_OPTS'] = '0'  # Disable oneDNN custom operations

# from flask import Flask, request, jsonify
# import io
# import PyPDF2
# import docx
# from transformers import pipeline

# app = Flask(__name__)

# def extract_text_from_pdf(file_object):
#     """
#     Extract text from a PDF file object.
#     """
#     try:
#         pdf_reader = PyPDF2.PdfReader(file_object)
#         text = ""
#         for page in pdf_reader.pages:
#             text += page.extract_text() + "\n"
#         return text
#     except Exception as e:
#         print(f"Error extracting text from PDF: {e}")
#         return ""

# def extract_text_from_docx(file_object):
#     """
#     Extract text from a DOCX file object.
#     """
#     try:
#         doc = docx.Document(file_object)
#         text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
#         return text
#     except Exception as e:
#         print(f"Error extracting text from DOCX: {e}")
#         return ""

# def summarize_text(full_text, max_length=500, min_length=200):
#     """
#     Summarize given text.
#     """
#     summarizer = pipeline("summarization", model="facebook/bart-large-cnn")
#     summary = summarizer(full_text, max_length=max_length, min_length=min_length, do_sample=False)
#     return summary[0]['summary_text']

# @app.route('/summarize_documents', methods=['POST'])
# def summarize_documents():
#     """
#     Handle document summarization.
#     """
#     if 'files' not in request.files:
#         return jsonify({'error': 'No files part in the request'}), 400
    
#     files = request.files.getlist('files')
    
#     if not files or files[0].filename == '':
#         return jsonify({'error': 'No files selected for uploading'}), 400
    
#     try:
#         full_text = ""
#         processed_files = []
        
#         for file in files:
#             file_object = io.BytesIO(file.read())
            
#             if file.filename.lower().endswith('.pdf'):
#                 text = extract_text_from_pdf(file_object)
#             elif file.filename.lower().endswith(('.docx', '.doc')):
#                 text = extract_text_from_docx(file_object)
#             else:
#                 continue
            
#             if text.strip():
#                 full_text += text + "\n\n---Document Separator---\n\n"
#                 processed_files.append(file.filename)
        
#         if not full_text.strip():
#             return jsonify({'error': 'No text could be extracted from the uploaded files'}), 400
        
#         summary = summarize_text(full_text)
        
#         return jsonify({
#             'summary': summary
#         })
    
#     except Exception as e:
#         return jsonify({'error': str(e)}), 500

# if __name__ == '__main__':
#     app.run(debug=False, port=3000, host='0.0.0.0')


import os
import re
import math
from flask import Flask, request, jsonify
from transformers import BartForConditionalGeneration, BartTokenizer
from PyPDF2 import PdfReader
from docx import Document
import torch

app = Flask(__name__)

# Load the BART model and tokenizer
def load_model():
    model_name = "facebook/bart-large-cnn"
    model = BartForConditionalGeneration.from_pretrained(model_name)
    tokenizer = BartTokenizer.from_pretrained(model_name)
    return model, tokenizer

model, tokenizer = load_model()

# Function to extract text from PDFs with chunk support for large documents
def extract_text_from_pdf(file, chunk_size=2000, overlap=500):
    reader = PdfReader(file)
    total_pages = len(reader.pages)
    full_text = ""
    
    for page in reader.pages:
        full_text += page.extract_text() + "\n"
    
    # Break large text into manageable chunks
    chunks = []
    words = full_text.split()
    for i in range(0, len(words), chunk_size - overlap):
        chunk = " ".join(words[i:i+chunk_size])
        chunks.append(chunk)
    
    return chunks, total_pages

# Function to extract text from Word documents
def extract_text_from_docx(file):
    doc = Document(file)
    text = "\n".join([para.text for para in doc.paragraphs if para.text])
    return [text]  # Return as a list to be consistent with PDF extraction

# Function to clean and preprocess text
def preprocess_text(text):
    # Remove special characters, extra whitespaces
    text = re.sub(r"[^\w\s.,!?]", "", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text

# Function to summarize text with dynamic length
def summarize_text(text, total_pages, min_length=None, max_length=None):
    # Dynamically adjust summary length based on document size
    if min_length is None:
        min_length = max(50, total_pages * 10)
    if max_length is None:
        max_length = max(200, total_pages * 30)
    
    inputs = tokenizer(text, max_length=1024, truncation=True, return_tensors="pt")
    
    summary_ids = model.generate(
        inputs.input_ids,
        min_length=min_length,
        max_length=max_length,
        length_penalty=2.0,
        num_beams=4,
        early_stopping=True
    )
    
    return tokenizer.decode(summary_ids[0], skip_special_tokens=True)

@app.route("/summarize", methods=["POST"])
def summarize():
    try:
        # Retrieve files from the request
        files = request.files.getlist("files")
        min_length = int(request.form.get("min_length", 0))
        max_length = int(request.form.get("max_length", 0))
        
        if not files:
            return jsonify({"error": "No files uploaded"}), 400
        
        all_chunks = []
        total_pages = 0
        
        for file in files:
            if file.filename.endswith(".pdf"):
                chunks, page_count = extract_text_from_pdf(file)
                total_pages += page_count
            elif file.filename.endswith(".docx"):
                chunks = extract_text_from_docx(file)
                total_pages = 1  # Assume single page for DOCX
            else:
                return jsonify({"error": f"Unsupported file type: {file.filename}"}), 400
            
            all_chunks.extend(chunks)
        
        # Preprocess and clean chunks
        cleaned_chunks = [preprocess_text(chunk) for chunk in all_chunks]
        
        # Combine chunks if they meet minimum length
        combined_text = " ".join([chunk for chunk in cleaned_chunks if len(chunk.split()) > 50])
        
        if len(combined_text.split()) > 50:
            # Generate summary
            summary = summarize_text(
                combined_text, 
                total_pages, 
                min_length if min_length > 0 else None, 
                max_length if max_length > 0 else None
            )
            
            return jsonify({
                "status": "success",
                "summary": summary,
                "total_pages": total_pages,
                "chunk_count": len(all_chunks)
            }), 200
        else:
            return jsonify({"error": "The text is too short for summarization"}), 400
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500


if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=3000)